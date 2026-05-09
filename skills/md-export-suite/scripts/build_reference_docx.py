#!/usr/bin/env python3
"""Build a clean Pandoc reference.docx for polished document exports."""

from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


STYLE_DIR = Path(__file__).resolve().parents[1] / "styles"


def load_style(style_id: str) -> dict:
    path = STYLE_DIR / f"{style_id}.json"
    if not path.exists():
        choices = ", ".join(sorted(p.stem for p in STYLE_DIR.glob("*.json") if p.name != "style_index.json"))
        raise SystemExit(f"Unknown style '{style_id}'. Available styles: {choices}")
    return json.loads(path.read_text(encoding="utf-8"))


def rgb(hex_value: str) -> RGBColor:
    value = hex_value.strip().lstrip("#")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def set_font(style, font_name: str, size: float, color: str, bold: bool = False) -> None:
    style.font.name = font_name
    style.font.size = Pt(size)
    style.font.color.rgb = rgb(color)
    style.font.bold = bold
    if style.element.rPr is not None:
        style.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_paragraph(style, before: float = 0, after: float = 6, line: float = 1.15) -> None:
    fmt = style.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def build(output: Path, company: str, style_id: str, style_token: str | None = None) -> None:
    style = load_style(style_id)
    colors = style["colors"]
    typography = style["typography"]
    document = style.get("document", {})
    ink = colors["ink"].strip().lstrip("#")
    body = colors["body"].strip().lstrip("#")
    muted = colors["muted"].strip().lstrip("#")
    primary = (style_token or colors["primary"]).strip().lstrip("#")
    display_font = typography.get("display", "Arial")
    body_font = typography.get("body", "Arial")
    mono_font = typography.get("mono", "Courier New")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    styles = doc.styles
    set_font(styles["Normal"], body_font, 10.8, body)
    set_paragraph(styles["Normal"], after=7, line=1.18)
    set_font(styles["Title"], display_font, float(document.get("title_size", 26)), ink, bold=True)
    set_paragraph(styles["Title"], after=14, line=1.02)

    heading_specs = [
        ("Heading 1", 19, primary, 18, 8),
        ("Heading 2", 15, ink, 14, 6),
        ("Heading 3", 12.5, ink, 10, 4),
        ("Heading 4", 11.5, muted, 8, 3),
    ]
    for name, size, color, before, after in heading_specs:
        set_font(styles[name], display_font, size, color, bold=True)
        set_paragraph(styles[name], before=before, after=after, line=1.08)

    for name in ("Source Code", "Code", "Verbatim Char"):
        if name in styles:
            set_font(styles[name], mono_font, 9.5, ink)
            set_paragraph(styles[name], before=4, after=7, line=1.0)

    header = section.header.paragraphs[0]
    header.text = company.upper()
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if header.runs:
        header.runs[0].font.name = body_font
        header.runs[0].font.size = Pt(8)
        header.runs[0].font.color.rgb = rgb(muted)
        header.runs[0].font.bold = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Page ")
    add_page_number(footer)
    for run in footer.runs:
        run.font.name = body_font
        run.font.size = Pt(8)
        run.font.color.rgb = rgb(muted)

    doc.add_paragraph("Reference Document", style="Title")
    doc.add_paragraph("This file defines default styles for Markdown exports.")
    doc.add_heading("Heading 1", level=1)
    doc.add_paragraph("Vietnamese text should remain readable: Trí tuệ nhân tạo, dữ liệu, quy trình.")
    doc.add_heading("Heading 2", level=2)
    doc.add_paragraph("Tables, lists, and code blocks inherit Word styles from this document.")

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> None:
    parser = argparse.ArgumentParser(description="Build a reference.docx for Pandoc.")
    parser.add_argument("--output", "-o", default="build/reference.docx", help="Output .docx path.")
    parser.add_argument("--company", default="MD Export Suite", help="Header label.")
    parser.add_argument("--style", default="institutional-clarity", help="Generic style ID.")
    parser.add_argument("--style-token", help="Optional primary accent hex override.")
    args = parser.parse_args()
    output = Path(args.output)
    build(output, args.company, args.style, args.style_token)
    print(output)


if __name__ == "__main__":
    main()
