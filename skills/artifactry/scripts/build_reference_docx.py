#!/usr/bin/env python3
"""Build a clean Pandoc reference.docx for polished document exports."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from style_resolver import load_style


def rgb(hex_value: str) -> RGBColor:
    value = hex_value.strip().lstrip("#")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def set_font(style, font_name: str, size: float, color: str, bold: bool = False) -> None:
    style.font.name = font_name
    style.font.size = Pt(size)
    style.font.color.rgb = rgb(color)
    style.font.bold = bold
    if style.element.rPr is not None:
        style.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_paragraph(style, before: float = 0, after: float = 6, line: float = 1.15) -> None:
    fmt = style.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_shading(style, fill: str) -> None:
    ppr = style.element.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:fill"), fill.strip().lstrip("#"))


def set_border_bottom(style, color: str, size: str = "8") -> None:
    ppr = style.element.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    bottom = pbdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        pbdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color.strip().lstrip("#"))


def configure_page(section, page_size: str) -> None:
    if page_size.lower() == "letter":
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
    else:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def build(output: Path, company: str, style_id: str, style_token: str | None = None, page_size: str = "a4") -> None:
    style, resolved_style = load_style(style_id)
    colors = style["colors"]
    typography = style["typography"]
    document = style.get("document", {})
    ink = colors["ink"].strip().lstrip("#")
    body = colors["body"].strip().lstrip("#")
    muted = colors["muted"].strip().lstrip("#")
    primary = (style_token or colors["primary"]).strip().lstrip("#")
    display_font = typography.get("display", "Arial")
    body_font = typography.get("body", "Arial")
    mono_font = typography.get("mono", "Courier New")

    doc = Document()
    section = doc.sections[0]
    configure_page(section, page_size)
    margins = document.get("margins_cm", [2.2, 2.0, 2.2, 2.2])
    section.top_margin = Cm(float(margins[0]))
    section.bottom_margin = Cm(float(margins[1]))
    section.left_margin = Cm(float(margins[2]))
    section.right_margin = Cm(float(margins[3]))

    styles = doc.styles
    set_font(styles["Normal"], body_font, 10.8, body)
    set_paragraph(styles["Normal"], after=7, line=1.18)
    set_font(styles["Title"], display_font, float(document.get("title_size", 26)), ink, bold=True)
    set_paragraph(styles["Title"], after=14, line=1.02)

    heading_specs = [
        ("Heading 1", 19, primary, 18, 8),
        ("Heading 2", 15, ink, 14, 6),
        ("Heading 3", 12.5, ink, 10, 4),
        ("Heading 4", 11.5, muted, 8, 3),
    ]
    for name, size, color, before, after in heading_specs:
        set_font(styles[name], display_font, size, color, bold=True)
        set_paragraph(styles[name], before=before, after=after, line=1.08)
        if name == "Heading 1":
            set_border_bottom(styles[name], primary, "10")

    for name in ("Source Code", "Code", "Verbatim Char"):
        if name in styles:
            set_font(styles[name], mono_font, 9.5, ink)
            set_paragraph(styles[name], before=4, after=7, line=1.0)
            if name == "Source Code":
                set_shading(styles[name], colors["surface"])

    for name in ("Block Text", "Quote"):
        if name in styles:
            set_font(styles[name], body_font, 10.8, body)
            set_paragraph(styles[name], before=8, after=8, line=1.18)
            set_shading(styles[name], colors["surface"])

    for name in ("Caption", "Image Caption", "Table Caption"):
        if name in styles:
            set_font(styles[name], body_font, 8.8, muted, bold=False)
            set_paragraph(styles[name], before=4, after=8, line=1.1)

    for name in ("List Bullet", "List Number", "Compact"):
        if name in styles:
            set_font(styles[name], body_font, 10.6, body)
            set_paragraph(styles[name], after=4, line=1.15)

    for table_name in ("Table Grid", "Light Shading Accent 1", "Medium Shading 1 Accent 1"):
        if table_name in styles:
            styles[table_name].font.name = body_font
            styles[table_name].font.size = Pt(9.5)

    header = section.header.paragraphs[0]
    header.text = f"{company.upper()} · {resolved_style.upper()}"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if header.runs:
        header.runs[0].font.name = body_font
        header.runs[0].font.size = Pt(8)
        header.runs[0].font.color.rgb = rgb(muted)
        header.runs[0].font.bold = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Page ")
    add_page_number(footer)
    for run in footer.runs:
        run.font.name = body_font
        run.font.size = Pt(8)
        run.font.color.rgb = rgb(muted)

    doc.add_paragraph("Reference Document", style="Title")
    doc.add_paragraph("This file defines default styles for Markdown exports.")
    doc.add_heading("Heading 1", level=1)
    doc.add_paragraph("Vietnamese text should remain readable: Trí tuệ nhân tạo, dữ liệu, quy trình.")
    doc.add_heading("Heading 2", level=2)
    doc.add_paragraph("Tables, lists, and code blocks inherit Word styles from this document.")
    quote_style = "Block Text" if "Block Text" in styles else "Normal"
    doc.add_paragraph("This is a block quote / callout surface.", style=quote_style)
    table = doc.add_table(rows=2, cols=3)
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, ["Route", "Output", "Status"]):
        cell.text = text
    for cell, text in zip(table.rows[1].cells, ["Markdown", "DOCX/PDF", "Validated"]):
        cell.text = text

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> None:
    parser = argparse.ArgumentParser(description="Build a reference.docx for Pandoc.")
    parser.add_argument("--output", "-o", default="build/reference.docx", help="Output .docx path.")
    parser.add_argument("--company", default="Artifactry", help="Header label.")
    parser.add_argument("--style", default="institutional-clarity", help="Generic style ID.")
    parser.add_argument("--style-token", help="Optional primary accent hex override.")
    parser.add_argument("--page-size", default="a4", choices=["a4", "letter"], help="Document page size.")
    args = parser.parse_args()
    output = Path(args.output)
    build(output, args.company, args.style, args.style_token, args.page_size)
    print(output)


if __name__ == "__main__":
    main()
